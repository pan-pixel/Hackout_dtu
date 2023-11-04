import os
from decouple import config


os.environ['OPENAI_API_KEY'] = config('OPENAI_API_KEY')
os.environ["SERPAPI_API_KEY"] = config('SERPAPI_API_KEY')

from gpt_index import SimpleDirectoryReader, GPTListIndex, GPTSimpleVectorIndex, LLMPredictor, PromptHelper
from langchain import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain.memory import ConversationBufferMemory
from langchain.agents import Tool, load_tools
from langchain.agents import initialize_agent
from langchain.tools import BaseTool
from bs4 import BeautifulSoup
import requests
from langchain.prompts.chat import SystemMessagePromptTemplate
from langchain import PromptTemplate
import csv
import docx
from docx.shared import Cm
from docx.shared import Pt
import re


llm = ChatOpenAI(
    temperature = 0,
    model_name = 'gpt-3.5-turbo'
)


tools = load_tools(["serpapi","llm-math","wikipedia"],llm=llm)


# custom tool for format
def chatbot(input_text):
    index = GPTSimpleVectorIndex.load_from_disk('index.json')
    response = index.query(input_text, response_mode="compact")
    return response

format_tool = Tool(
    name = 'Equity Report Template',
    func = chatbot,
    description = "Useful for when you need to create equity research report, consult the format from this. Input should be about the format or essentials of equity research report"
)
# tool for stock information
class WebPageTool(BaseTool):
    name = "Get Webpage"
    description = "Useful for when you need to get the content from a specific webpage"

    def _run(self, webpage: str):
        response = requests.get(webpage)
        html_content = response.text

        def strip_html_tags(html_content):
            soup = BeautifulSoup(html_content, "html.parser")
            stripped_text = soup.get_text()
            return stripped_text

        stripped_content = strip_html_tags(html_content)
        if len(stripped_content) > 4000:
            stripped_content = stripped_content[:4000]
        return stripped_content
    
    def _arun(self, webpage: str):
        raise NotImplementedError("This tool does not support async")

page_getter = WebPageTool()
tools.append(page_getter)
tools.append(format_tool)


memory = ConversationBufferWindowMemory(
    memory_key='chat_history',
    k=3,
    return_messages=True
)

company = input("Enter company name : ")
data = company
company = company.replace(" ", "+")
url = "https://www.google.com/search?q=" + company + "+stock+business+today"
r = requests.get(url=url)
htmlContent = r.content
soup = BeautifulSoup(htmlContent, "html.parser")
links = soup.find_all('a', href=True)
for link in links:
    if 'www.businesstoday.in' in link['href']:
        final_link = link['href'].split("&")[0][7:]


print(final_link)

# memory = ConversationBufferMemory(memory_key="chat_history")
agent = initialize_agent(tools= tools, llm = llm, agent ='chat-conversational-react-description', verbose = True, max_iterations=5,
    early_stopping_method='generate',
    memory=memory)

fixed_prompt = '''Assistant is a large language model trained by OpenAI.

Assistant is designed to be able to assist with a wide range of tasks, from answering simple questions to providing in-depth explanations and discussions on a wide range of topics. As a language model, Assistant is able to generate human-like text based on the input it receives, allowing it to engage in natural-sounding conversations and provide responses that are coherent and relevant to the topic at hand.

Assistant doesn't know anything about company's information and should use a tool for questions about these topics.

Assistant should provide the data in maximum detail possible

Assistant is constantly learning and improving, and its capabilities are constantly evolving. It is able to process and understand large amounts of text, and can use this knowledge to provide accurate and informative responses to a wide range of questions. Additionally, Assistant is able to generate its own text based on the input it receives, allowing it to engage in discussions and provide explanations and descriptions on a wide range of topics.

Overall, Assistant is a powerful system that can help with a wide range of tasks and provide valuable insights and information on a wide range of topics. Whether you need help with a specific question or just want to have a conversation about a particular topic, Assistant is here to assist.'''


agent.agent.llm_chain.prompt.messages[0].prompt.template = fixed_prompt

company_info = agent.run(f"What is {company} and what does it do?")


business_des = agent.run(f"What is the history of {company} on wikipedia?")


pe = agent.run(f"what is the Price-to-Earnings of {data} on {final_link}")
share_price = agent.run(f"what is the stock price of {data} on {final_link}")
expected_earning = agent.run(f"what is the return on equity of {data} on {final_link} of 2023")



def find_numerical_value(string):
    pattern = r"[-+]?\d*\.?\d+|\d+"  # Regular expression pattern to match numbers with optional sign
    matches = re.findall(pattern, string)
    numerical_values = [float(match) for match in matches]
    return numerical_values

numerical_value = find_numerical_value(pe)
pe = numerical_value[0]

numerical_value = find_numerical_value(share_price)
share_price = numerical_value[0]

numerical_value = find_numerical_value(expected_earning)
expected_earning = numerical_value[0]


forward_pe =share_price/expected_earning
price_target = pe/forward_pe
# print(pe)
# pe = pe.split("₹")
# pe = float(pe[1])
# share_price = share_price.split("₹")
# expected_earning = expected_earning.split("₹")
# share_price = (float(share_price[1]))
# expected_earning = (float(expected_earning[1]))

# forward_pe = share_price/expected_earning
# price_target = pe/forward_pe


# f.writelines(price_target)

result = agent.run(f"What is the current status of the {company}?")



invest_summ = agent.run(f"summarize the data in bullet form of {data} on {final_link}")


compi_posi = agent.run(f"who are the competitors of {data} on {final_link}")


value = agent.run(f"summarize all the ration analysis in bullet form of {data} on {final_link}")






topic=data #company name
Curent_price=share_price #curent price
target_price=price_target #target price
if target_price > share_price:

    recommendation="Buy"
else:
    recommendation="Sell"

company_info= company_info
buisness_des= business_des
invest_summ=invest_summ
compi_posi=compi_posi
valu=value


doc = docx.Document()
section = doc.sections[0]
section.page_width = docx.shared.Inches(8.27)  # 21cm
section.page_height = docx.shared.Inches(11.69)  # 29.7cm
sections = doc.sections
margins=0.4   #Page Margin (in inche)
for section in sections:
    section.top_margin = docx.shared.Inches(margins)
    section.bottom_margin = docx.shared.Inches(margins)
    section.left_margin = docx.shared.Inches(margins)
    section.right_margin = docx.shared.Inches(margins)

heading = doc.add_heading(topic.upper(), level=1)
heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading.style.font.size = Pt(28)
heading.style.font.name = 'Times New Roman'

table = doc.add_table(rows=2, cols=2)
row = table.rows[0]
cell1 = row.cells[0]
cell1.text = "Current Price: " + str(Curent_price)
cell2 = row.cells[1]
cell2.text = "Recommendation By : "

row = table.rows[1]
cell1 = row.cells[0]
cell1.text = "Target Price: " + target_price
cell2 = row.cells[1]
cell2.text = "P/E:" + recommendation

title = doc.add_heading(level=4)
title.text = "Document Title"

paragraph = doc.add_paragraph(company_info)
font = paragraph.style.font           
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_heading(level=4)
title.text = "Buisness Description"

paragraph = doc.add_paragraph(buisness_des)
font = paragraph.style.font          
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_heading(level=4)
title.text = "Investment Summary"

paragraph = doc.add_paragraph(invest_summ)
font = paragraph.style.font          
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_heading(level=4)
title.text = "Competetive Positioning"

paragraph = doc.add_paragraph(compi_posi)
font = paragraph.style.font          
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_heading(level=4)
title.text = "Valuation"

paragraph = doc.add_paragraph(valu)
font = paragraph.style.font          
font.name = 'Times New Roman'
font.size = Pt(12)

print("report generated")

doc.save(topic+".docx")



# print(result)