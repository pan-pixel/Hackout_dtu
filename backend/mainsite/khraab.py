from django.shortcuts import render, redirect
from time import sleep
import os
from decouple import config
import openai
import requests
import json
import sys
import requests 

import nltk
from textblob import TextBlob 
from newspaper import Article 
from .models import News_analysis





os.environ['OPENAI_API_KEY'] = config('OPENAI_API_KEY')
os.environ["SERPAPI_API_KEY"] = config('SERPAPI_API_KEY')

openai.api_key = config('OPENAI_API_KEY')

# openai.api_key = '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'

from gpt_index import SimpleDirectoryReader, GPTListIndex, GPTSimpleVectorIndex, LLMPredictor, PromptHelper
from langchain import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.chains.conversation.memory import ConversationBufferWindowMemory
from langchain.memory import ConversationBufferMemory
from langchain.agents import Tool, load_tools
from langchain.agents import initialize_agent
from langchain.tools import BaseTool
from bs4 import BeautifulSoup
import requests
from langchain.prompts.chat import SystemMessagePromptTemplate
from langchain import PromptTemplate
import csv
import docx
from docx.shared import Cm
from docx.shared import Pt
import re
from docx.oxml import OxmlElement
from docx.shared import Inches
import matplotlib.pyplot as plt
import yfinance as yf
from io import BytesIO
from docx2pdf import convert



def clean():
    # Define the pattern to search for "action_input": "..."
    pattern = r'"action_input":\s+"(.*?)"'

    # Read the text file
    with open('/Users/prathamarora/Documents/glb_hack/backend/output.txt', 'r', encoding='utf-8') as file:
        text = file.read()

    # Find all matches in the text
    matches = re.finditer(pattern, text)

    # Extract and print data from all matches
    final_text = ""
    for match in matches:
        action_input_data = match.group(1)
        if "https://" not in action_input_data:
            final_text += action_input_data

    return final_text

def gpt_format(company):
    info = clean()
    # prompt = "I am giving you basic details and current stock and investment information about"+company+", format it quite well with written content separate and all the statics properly laid out and each of them also explained in detail with respect to company.For the written material make sure to elaborate about the company as much as you can, you are free to also add initaitves and other details about the company which you think are impacting its performance in long term.  You are completely free to use language of your own and hence make it quite elaborated and professionaly written as it is to be used as content for professional equity investment report, write a minimum of three pages as per the above instructions and make sure to add headings. Here are the details - " + info
    # prompt = "I would like you to provide a well-structured and comprehensive report on Tata Consultancy Services (TCS), encompassing essential details, current stock and investment information, along with in-depth explanations of each statistic in the context of the company. List all the numeric values together as well. In addition, please elaborate on the company's background and operations, and feel free to include information about key initiatives and factors that may influence its long-term performance. You have the creative freedom to use professional language and ensure the report is elaborative and suitable for a professional equity investment report. Your report should be at least three pages in length, and don't forget to include clear headings for each section, amd before and after each heading add '&' for clearity. Here are the details - " + info
    prompt = "I would like you to provide a well-structured and comprehensive report on " + company + ", encompassing essential details, current stock and investment information, along with in-depth explanations of each statistic in the context of the company. List all the numeric values together as well. In addition, please elaborate on the company's background and operations, and feel free to include information about key initiatives and factors that may influence its long-term performance. You have the creative freedom to use professional language and ensure the report is elaborative and suitable for a professional equity investment report. Your report should be at least three pages in length, and don't forget to include clear headings for each section, to mark starting of the heading use '~' Here are the details - " + info + "Make sure to add the stock symbol at the end, the output of stock symbol shall be like ~stock_symbol, that's it nothing else"
    response = openai.ChatCompletion.create(
        model = 'gpt-3.5-turbo',
        messages=[
            {"role": "user", "content": prompt}
        ]

    )
    content = response['choices'][0]['message']['content']
    splitData = content.split("~")
    for c in splitData:
        c = c.split("\n", 1)

    final = splitData[1:]
    return final
    

llm = ChatOpenAI(
    temperature = 0,
    model_name = 'gpt-3.5-turbo'
)


tools = load_tools(["serpapi","llm-math","wikipedia"],llm=llm)

# custom tool for format
def chatbot(input_text):
    index = GPTSimpleVectorIndex.load_from_disk('index.json')
    response = index.query(input_text, response_mode="compact")
    return response

format_tool = Tool(
    name = 'Equity Report Template',
    func = chatbot,
    description = "Useful for when you need to create equity research report, consult the format from this. Input should be about the format or essentials of equity research report"
)
# tool for stock information
class WebPageTool(BaseTool):
    name = "Get Webpage"
    description = "Useful for when you need to get the content from a specific webpage"

    def _run(self, webpage: str):
        response = requests.get(webpage)
        html_content = response.text

        def strip_html_tags(html_content):
            soup = BeautifulSoup(html_content, "html.parser")
            stripped_text = soup.get_text()
            return stripped_text

        stripped_content = strip_html_tags(html_content)
        if len(stripped_content) > 4000:
            stripped_content = stripped_content[:4000]
        return stripped_content
    
    def _arun(self, webpage: str):
        raise NotImplementedError("This tool does not support async")



def report_generator(company):
    output_file = open("output.txt","w")
    sys.stdout = output_file

    page_getter = WebPageTool()
    tools.append(page_getter)
    tools.append(format_tool)


    memory = ConversationBufferWindowMemory(
        memory_key='chat_history',
        k=3,
        return_messages=True
    )
    data = company
    company = company.replace(" ", "+")
    url = "https://www.google.com/search?q=" + company + "+stock+business+today"
    r = requests.get(url=url)
    htmlContent = r.content
    soup = BeautifulSoup(htmlContent, "html.parser")
    links = soup.find_all('a', href=True)
    for link in links:
        if 'www.businesstoday.in' in link['href']:
            final_link = link['href'].split("&")[0][7:]


    # print(final_link)

    # memory = ConversationBufferMemory(memory_key="chat_history")

    agent = initialize_agent(tools= tools, llm = llm, agent ='chat-conversational-react-description', verbose = True, max_iterations=5,
        early_stopping_method='generate',
        memory=memory)

    fixed_prompt = '''Assistant is a large language model trained by OpenAI.

    Assistant is designed to be able to assist with a wide range of tasks, from answering simple questions to providing in-depth explanations and discussions on a wide range of topics. As a language model, Assistant is able to generate human-like text based on the input it receives, allowing it to engage in natural-sounding conversations and provide responses that are coherent and relevant to the topic at hand.

    Assistant doesn't know anything about company's information and should use a tool for questions about these topics.

    Assistant should provide the data in maximum detail possible

    Assistant is constantly learning and improving, and its capabilities are constantly evolving. It is able to process and understand large amounts of text, and can use this knowledge to provide accurate and informative responses to a wide range of questions. Additionally, Assistant is able to generate its own text based on the input it receives, allowing it to engage in discussions and provide explanations and descriptions on a wide range of topics.

    Overall, Assistant is a powerful system that can help with a wide range of tasks and provide valuable insights and information on a wide range of topics. Whether you need help with a specific question or just want to have a conversation about a particular topic, Assistant is here to assist.'''


    agent.agent.llm_chain.prompt.messages[0].prompt.template = fixed_prompt

    company_info = agent.run(f"What is {company} and what does it do?")


    business_des = agent.run(f"What is the history of {company} on wikipedia?")


    pe = agent.run(f"what is the Price-to-Earnings of {data} on {final_link}")
    share_price = agent.run(f"what is the stock price of {data} on {final_link}")
    expected_earning = agent.run(f"what is the return on equity of {data} on {final_link} of 2023")




    result = agent.run(f"What is the current status of the {company}?")



    invest_summ = agent.run(f"summarize the data in bullet form of {data} on {final_link}")


    compi_posi = agent.run(f"who are the competitors of {data} on {final_link}")


    value = agent.run(f"summarize all the ration analysis in bullet form of {data} on {final_link}")

    # ballance_sheet = agent.run(f"Find the balance sheet of the {company} and provide all its data")

    output_file.close()


def initial_setup(doc, topic):
    section = doc.sections[0]
    section.page_width = docx.shared.Inches(8.27)  # 21cm
    section.page_height = docx.shared.Inches(11.69)  # 29.7cm
    sections = doc.sections
    margins=0  #Page Margin (in inche)
    for section in sections:
        section.top_margin = docx.shared.Inches(margins)
        section.bottom_margin = docx.shared.Inches(margins)
        section.left_margin = docx.shared.Inches(margins)
        section.right_margin = docx.shared.Inches(margins)

    paragraph = doc.add_paragraph()
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    img_path = '/Users/prathamarora/Documents/glb_hack/backend/static/assests/header.png' 
    run = paragraph.add_run()
    picture = run.add_picture(img_path)             #width=docx.shared.Inches(9.0)

    font_size = docx.shared.Pt(11)
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0 #Line-Spacing

    heading = doc.add_heading(topic.upper(), level=1)
    heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    heading.style.font.size = Pt(28)
    heading.style.font.name = 'Times New Roman'

def main_content(doc, data, company):
    # data = ['Introduction', "\nThis equity investment report provides an overview of Tata Consultancy Services (TCS), a leading Indian multinational IT services and consulting company. The report covers essential details about the company, including its background, operations, and key initiatives. It also includes current stock and investment information, along with an in-depth explanation of each statistic in the context of TCS. This report aims to provide investors with comprehensive insights into TCS's performance and potential for long-term growth.\n\n", 'Company Background and Operations', '\nTCS, a part of the Tata Group, is a major player in the global IT services industry. It operates in 150 locations across 46 countries, making it a truly multinational company. TCS offers a wide range of IT services, including software development, consulting, and business process outsourcing. With over 600,000 employees worldwide, TCS has established itself as a leader in the industry.\n\n', 'Key Initiatives and Factors Influencing Long-term Performance', "\nTCS has undertaken several key initiatives to drive its long-term performance. These include investing in emerging technologies like artificial intelligence (AI), cloud computing, and blockchain. By focusing on these areas, TCS aims to stay ahead of the curve and address evolving client needs. Additionally, TCS prioritizes innovation and R&D efforts to develop cutting-edge solutions and enhance its competitive advantage.\n\nThe company's strong client relationships and robust delivery capabilities have played a significant role in its long-term success. TCS has a diverse client base across industries, including banking, healthcare, retail, and telecommunications. This diversity helps TCS mitigate the risks associated with industry-specific fluctuations and ensures a more stable revenue stream.\n\nTCS also emphasizes skill development and talent retention as crucial factors for sustained growth. The company invests heavily in employee training and development programs to enhance its workforce's capabilities. TCS's dedication to talent management enables it to deliver high-quality services, gain client trust, and attract and retain top-tier talent.\n\nThe company's long-standing commitment to corporate social responsibility (CSR) is another aspect contributing to its long-term performance. TCS actively engages in various CSR initiatives, including education, environment, and community development programs. These initiatives not only demonstrate TCS's commitment to social welfare but also contribute to its reputation and brand image.\n\n", 'Current Stock and Investment Information', "\nTCS's stock performance is an important consideration for equity investors. As of the provided webpage, the stock price of TCS is ₹3,350.90, exhibiting a consistent upward trend over the years. The company's market capitalization is ₹12,26,111.52 Cr., making it the second-largest Indian company.\n\nKey statistical information related to TCS's performance is as follows:\n\n1. Price-to-Earnings (P/E) Ratio: The P/E ratio of TCS is 27.46. This ratio indicates the market's valuation of the company relative to its earnings. It is commonly used to assess whether a stock is overvalued or undervalued.\n\n2. Earnings Per Share (EPS): TCS has an EPS of ₹122.04, which indicates the company's profitability on a per-share basis. Investors often consider EPS to evaluate a company's financial health and growth potential.\n\n3. Dividend Yield: The dividend yield of TCS is 3.43%. This metric reflects the return on investment from dividend payments, which is particularly important for income-focused investors.\n\n4. Price-to-Book (P/B) Ratio: TCS's P/B ratio is 12.18. The P/B ratio compares a company's market value to its book value and can help investors assess whether a stock is trading at a fair value.\n\n5. Beta: TCS's beta is 0.97, indicating a relatively low level of systematic risk. This suggests that TCS's stock price is less volatile than the market average.\n\n", 'Competitors', "\nTCS faces competition from several prominent players in the IT services industry. The competitors listed on the provided webpage include Infosys, HCL Technologies, L&T Technology Services, Sonata Software, Birlasoft, Mphasis, and Wipro. It is important to consider the competitive landscape when assessing TCS's performance and potential growth prospects.\n\n", 'Conclusion', "\nTCS, as a leading IT services and consulting company, has exhibited strong performance and positioned itself as a global leader in the industry. The company's diversified client base, focus on emerging technologies, and commitment to talent development and CSR contribute to its long-term success. Investors should examine the provided stock and investment information, along with the key initiatives and factors influencing TCS's performance, to assess its investment potential. However, it is crucial to conduct further research and due diligence before making investment decisions.", 'TCS']
    data = data
    global symbol
    prompt = "I need stock symbol of "+company+". The output format must be followed and is as follows - [stock_symbol]" 
    response = openai.ChatCompletion.create(
        model = 'gpt-3.5-turbo',
        messages=[
            {"role": "user", "content": prompt}
        ]
    )
    symbol = response['choices'][0]['message']['content']

    if"[" in symbol:
        st = symbol.index("[")
        en = symbol.index("]")
        symbol = symbol[st+1:en] + ".NS"
    else:
        symbol = symbol.split(" ")[-1] + ".NS"
    
    data = [text.replace("\n\n", "\n") for text in data]
    graphs_list = ["graph1", "graph2"]
    function_mapping = {
        "graph1" : graph1,
        "graph2" : graph2
    }
    for i in range(len(data)):
        
        if i%2 == 0:
            title = doc.add_heading(level=4)
            title.paragraph_format.left_indent = Pt(24)
            title.paragraph_format.right_indent = Pt(24)
            title.text = data[i]
            
        else:
            paragraph = doc.add_paragraph(data[i])
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.right_indent = Pt(24)
            paragraph.paragraph_format.line_spacing = Inches(0.3)
            paragraph.space_after = Pt(8)
            font = paragraph.style.font           
            font.name = 'Times New Roman'
            font.size = Pt(12)

        # if i>3 and i<3+len(graphs_list)+1:
        #         function = function_mapping[graphs_list[i-4]]
        #         function(symbol, doc)



def about(request):
    return render(request,"about.html")



# Create your views here.
def home(request):
    # sleep(1)
    if request.method == 'POST':
        company = request.POST['company']
        report_generator(company)
        ans = gpt_format(company)
        doc = docx.Document()
        topic = "report"
        symbol = ""
        initial_setup(doc, company)
        main_content(doc, ans, company)

        doc.save(topic+".docx")
        instructions = {
            'parts': [
            {
            'file': 'document'
            }
            ]
            }

        response = requests.request(
        'POST',
        'https://api.pspdfkit.com/build',
        headers = {
        'Authorization': 'Bearer pdf_live_LaJ5ZY4Auuz94CJ5Ch8upK0LVOzi3IPPX3T7GPB5QHX'
        },
        files = {
            'document': open('/Users/prathamarora/Documents/glb_hack/backend/report.docx', 'rb')
        },
        data = {
            'instructions': json.dumps(instructions)
        },
        stream = True
        )

        if response.ok:
            with open('/Users/prathamarora/Documents/glb_hack/backend/static/assests/report.pdf', 'wb') as fd:
                for chunk in response.iter_content(chunk_size=8096):
                    fd.write(chunk)
        else:
            print(response.text)
            exit()


        
        url = 'https://newsapi.org/v2/everything?'

        parameters = {
    'q': company, # query phrase
    'pageSize': 2,  # maximum is 100
    # 'category': 'business',
    'source':'moneycontrol.com',
    'language':'en',
    # 'country':'in',
    'apiKey':  'f4fe8e6439a7460d85cdc8a5491f63c1',
    'sortBy':'relevancy',
}


        member = News_analysis.objects.all()
        member.delete()
        response = requests.get(url, params=parameters)
        response_json = response.json()

        count = 0
        for i in response_json['articles']:
            first_link = (i['url'])
            headline = (i['title'])
            mood = ""
            if(count>3):
                break
            count+=1

            # print(first_link)

            article = Article(first_link)
            article.download() # Downloads Article
            article.parse() # Removes HTML and other elements 
            article.nlp() #Prepares it for natural language processing 
            text =article.summary # creates a summary of the article 
            blob= TextBlob(text)  # Textblob object is created for nlp
            sentiment=blob.sentiment.polarity # Finds the sentiment between -1 and 1 
            if sentiment >=0.25:
                # print("Good Sentiment around company , You can invest")
                mood = "Positive"
            elif sentiment >=0:
                # print("Neutral Sentiment")
                mood = "Neutral"

            else: 
                # print("Bad news for the company")
                mood = "Negative"

            # p = News_analysis.objects.create(company = company, headline=headline, link = first_link, sentiment=mood)
            new_file = News_analysis(company = company, headline=headline, link = first_link, sentiment=mood)
        
            new_file.save()

        return redirect('search')
    return render(request,"main.html")


def search(request):


    news_ana = News_analysis.objects.all()
    if request.method == 'POST':
        company = request.POST['company']
        print(company)
        report_generator(company)
        ans = gpt_format(company)
        doc = docx.Document()
        topic = "report"
        symbol = ""
        initial_setup(doc, company)
        main_content(doc, ans, company)

        doc.save(topic+".docx")
        instructions = {
            'parts': [
            {
            'file': 'document'
            }
            ]
            }

        response = requests.request(
        'POST',
        'https://api.pspdfkit.com/build',
        headers = {
        'Authorization': 'Bearer pdf_live_LaJ5ZY4Auuz94CJ5Ch8upK0LVOzi3IPPX3T7GPB5QHX'
        },
        files = {
            'document': open('/Users/prathamarora/Documents/glb_hack/backend/report.docx', 'rb')
        },
        data = {
            'instructions': json.dumps(instructions)
        },
        stream = True
        )

        if response.ok:
            with open('/Users/prathamarora/Documents/glb_hack/backend/static/assests/report.pdf', 'wb') as fd:
                for chunk in response.iter_content(chunk_size=8096):
                    fd.write(chunk)
        else:
            print(response.text)
            exit()



        # url = 'https://newsapi.org/v2/everything?'

        # parameters = {
        #     'q': company, # query phrase
        #     'pageSize': 2,  # maximum is 100
        #     'language':'en',
        #     'sources':'moneycontrol.com',
        #     # 'country':'in',
        #     'apiKey':  'f4fe8e6439a7460d85cdc8a5491f63c1',
        #     'sortBy':'relevancy',
        # }



        # response = requests.get(url, params=parameters)
        # response_json = response.json()

        # count = 0

        # for i in response_json['articles']:
        #     first_link = (i['url'])
        #     headline = (i['title'])
        #     mood = ""
        #     if(count>2):
        #         break
        #     count+=1

        #     # print(first_link)

        #     article = Article(first_link)
        #     article.download() # Downloads Article
        #     article.parse() # Removes HTML and other elements 
        #     article.nlp() #Prepares it for natural language processing 
        #     text =article.summary # creates a summary of the article 
        #     blob= TextBlob(text)  # Textblob object is created for nlp
        #     sentiment=blob.sentiment.polarity # Finds the sentiment between -1 and 1 
        #     if sentiment >=0.25:
        #         # print("Good Sentiment around company , You can invest")
        #         mood = "Positive"
        #     elif sentiment >=0:
        #         # print("Neutral Sentiment")
        #         mood = "Neutral"

        #     else: 
        #         # print("Bad news for the company")
        #         mood = "Negative"

        #     # p = News_analysis.objects.create(company = company, headline=headline, link = first_link, sentiment=mood)
        #     new_file = News_analysis(company = company, headline=headline, link = first_link, sentiment=mood)
        
        #     new_file.save()

            
        # context = {
        #     headline:headline,
        #     link:first_link,
        #     sentiment:mood,
        # }

        return redirect('search')
    
    data = {
        'news_ana':news_ana
    }
    return render(request,"aftersearch.html",data)





def gameLearn(request):
    

    return render(request,"learn.html")